from docx import Document
from docx.shared import Inches

# Create document
doc = Document()
doc.add_heading('Agent Security Audit Report', level=1)

# Add table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Vulnerability'
hdr_cells[1].text = 'Risk'
hdr_cells[2].text = 'Severity'
hdr_cells[3].text = 'Action Item'

# Add data rows
for row in data:
    cells = table.add_row().cells
    cells[0].text = row["Vulnerability"]
    cells[1].text = row["Risk"]
    cells[2].text = row["Severity"]
    cells[3].text = row["Action Item"]

# Save to file
file_path = "Agent_Security_Audit_Report.docx"
doc.save(file_path)


